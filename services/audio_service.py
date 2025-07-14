import os
import re #regax
import json
import shutil
import uuid
import requests
import math
import base64
import asyncio
import hashlib
import tempfile
import traceback

from utils import *
from utils.prompt_utils import get_prompt_by_category
from uuid import uuid4
from typing import List,Optional
from langchain_anthropic import ChatAnthropic
from langchain_openai import ChatOpenAI
from langchain_core.prompts import PromptTemplate
from langchain.chains import LLMChain
from langchain_core.documents import Document as LangchainDocument
# from langchain.chains.combine_documents.stuff import StuffDocumentsChain
# from langchain_community.document_loaders import UnstructuredFileLoader
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_unstructured import UnstructuredLoader
from fastapi import UploadFile, HTTPException
from moviepy import VideoFileClip, AudioFileClip
from pydub import AudioSegment
from docx.api import Document as DocxDocument
from dbcontroller.supabase import SupabaseConnector
# from dbcontroller.mongo import DBConnector

from warnings import filterwarnings
filterwarnings("ignore")

#mongo = DBConnector()
supabase = SupabaseConnector()

# Audio_segment length(For transcribe)
SEGMENT_LENGTH = 10  # Initially, 30

# LLM model
anthropic_model = ChatAnthropic(
    model=LLM_MODEL,
    api_key=LLM_KEY,
    max_tokens=4096
)

openai_model = ChatOpenAI(
    model="gpt-4o-mini",         # ปรับเป็นโมเดลที่ต้องการใช้งาน 
    openai_api_key=GPT_TOKEN,
    max_tokens=4096              # จำนวนโทเค็นสูงสุดในแต่ละครั้ง
)

# Prompt and LLM chain
prompt_path = os.path.join("prompts", "LLMprompt_task3.txt")
llmPrompt = PromptTemplate.from_file(prompt_path)
# llm_chain = LLMChain(llm=anthropic_model, prompt=llmPrompt)
# stuff_chain = StuffDocumentsChain(llm_chain=llm_chain,
#                                   document_variable_name="text_input_here")
# stuff_chain = create_stuff_documents_chain(anthropic_model, llmPrompt,
#                                     document_variable_name="text_input_here",)
stuff_chain = create_stuff_documents_chain(openai_model, llmPrompt,
                                    document_variable_name="text_input_here",)

# TTS headers and payload template
headers = {
    'Botnoi-Token': BOTNOI_TOKEN,
    'Content-Type': 'application/json'
}

# Language has been removed --> "language": "th"
payload_template = {
    "text": "",
    "speaker": "",
    "volume": 1.5,
    "speed": 1.0,
    "type_media": "wav",
    "save_file": "true",
}

#------------------UPLOAD SCRIPT----------------#
def docx_to_txt(docx_path: str, txt_path: str):
    """
    Convert a .docx file into .txt and remove the .docx file

    Parameters:
    docx_path (str): system path to the .docx file
    txt_path (str): path to save the .txt file
    """
    document = DocxDocument(docx_path)
    f = open(txt_path, "w", encoding="utf-8")
    for p in document.paragraphs:
        if p.text != ".":
            f.write(p.text + "\n")
            f.write("\n")
    if os.path.exists(docx_path):
        os.remove(docx_path)

async def upload_script_svc(session_id: str, script: List[UploadFile]):
    """
    Uploads a script file, replaces existing script if found, and saves metadata in DB.

    Parameters:
    session_id (str): Unique identifier for this user.
    script (List[UploadFile]): Uploaded script files.
    """
    supabase.add_or_update_section(session_id, "processing", "start-upload_script_svc")
    print("[INFO]: Started uploading script")

    # old = mongo.find_script(session_id)
    old = supabase.find_script(session_id)
    if old:
        os.remove(os.path.join(SERVICE_DIR, SCRIPT_DIR, session_id, old['session_id'] + '.txt'))
        # mongo.del_script(session_id)
        supabase.del_script(session_id)
    filepath = os.path.join(SERVICE_DIR, SCRIPT_DIR, session_id, session_id + '.txt')
    script_sub = os.path.join(SERVICE_DIR, SCRIPT_DIR, session_id)
    if not os.path.exists(script_sub):
        os.makedirs(script_sub, exist_ok=True)

    ext = os.path.splitext(script[0].filename)[-1].lower()
    print("Extension:", ext)
    if ext=='.docx':
        docx_path = os.path.join(script_sub, session_id + '.docx')
        with open(docx_path, "wb") as buffer:
            shutil.copyfileobj(script[0].file, buffer)
        docx_to_txt(docx_path, filepath)
    elif ext=='.txt':
        with open(filepath, "wb") as buffer:
            shutil.copyfileobj(script[0].file, buffer)
    else:
        print("[ERROR]: file type not supported.")
        supabase.add_or_update_section(session_id, "fail", "error-upload_script_svc")
        return
    
    transcript_data = get_transcript_data(session_id=session_id)
    transcribe_text = ' '.join(transcript_data['data'])
    metadata = {
        "uid": session_id,
        "summed": False,
        "session_id": session_id,
        "transcribe_text": transcribe_text
    }

    # mongo.add_script(metadata)
    supabase.add_script(metadata)
    supabase.add_or_update_section(session_id, "processing", "end-upload_script_svc")
    return

## move to util
def md5_storage(file_path: str, chunk_size: int = 8192) -> str:
    """
    Calculate hash of a file in disk memory.

    Parameters:
    file_path (str): The system path to the file.
    chunk_size (int): Hash chunk size.

    Returns:
    str: Hash string.
    """
    hash_md5 = hashlib.md5()
    with open(file_path, "rb") as f:
        while chunk := f.read(chunk_size):
            hash_md5.update(chunk)
    return hash_md5.hexdigest()

## move to util
async def md5_upload(upload_file: UploadFile, chunk_size: int = 8192) -> str:
    """
    Calculate hash of an UploadFile.

    Parameters:
    upload_file (UploadFile): UploadFile from an API endpoint.
    chunk_size (int): Hash chunk size.

    Returns:
    str: Hash string.
    """
    hash_md5 = hashlib.md5()
    while True:
        chunk = await upload_file.read(chunk_size)
        if not chunk:
            break
        hash_md5.update(chunk)
    await upload_file.seek(0)  # Reset the file pointer after reading
    return hash_md5.hexdigest()

## move to util
# Function to compare MD5 of physical file and UploadFile
async def compare_md5(file_path: str, upload_file: UploadFile) -> bool:
    """
    Compare hash of static file and UploadFile.

    Parameters:
    file_path (str): The system path to the file.
    upload_file (UploadFile): UploadFile from an API endpoint.

    Returns:
    bool: True if the hashes are identical, False if not.
    """
    disk_md5 = md5_storage(file_path)
    upload_md5 = await md5_upload(upload_file)
    return disk_md5 == upload_md5

def is_audio_file(file_name):
    audio_extensions = ('.mp3', '.wav', '.flac', '.ogg', '.aac', '.m4a', '.wma')
    return file_name.lower().endswith(audio_extensions)

async def upload_audios_svc(session_id: str, audios: List[UploadFile]):
    # mongo.update_used(session_id)
    supabase.update_used(session_id)
    user_path = os.path.join(SERVICE_DIR, VID_SRC_DIR, session_id)
    if not os.path.exists(user_path):
        os.makedirs(user_path, exist_ok=True)

    vid_infos = []
    for audio in audios:
        name, ext = os.path.splitext(audio.filename)
        # Generate a random UID instead of using the filename
        random_uid = str(uuid4())
        name = random_uid
        # dupe = mongo.match_fname(session_id, name)
        dupe = supabase.match_fname(session_id,"".join([name, ext]) )
        ## change to uid.wav
        if dupe:
            i = 1
            dpath = os.path.join(user_path, f"{dupe[0]['filename']}")
            same = await compare_md5(file_path=dpath, upload_file=audio)
            # same = video_service.compare_md5(file_path=dpath, upload_file=audio)
            if same:
                continue
            else:
                # while len(mongo.match_fname(session_id, name + f"_({i}")) != 0:
                while len(supabase.match_fname(session_id, name + f"_({i})")) != 0:
                    i += 1
                #name = name + f"_({i})"

        if is_audio_file(audio.filename):
            audio_output_folder = os.path.join(SERVICE_DIR, CACHE_DIR, session_id, "audio_segments")
            if not os.path.exists(audio_output_folder):
                os.makedirs(audio_output_folder, exist_ok=True)

            # audio_path = os.path.join(audio_output_folder, f'{session_id}{ext}')
            audio_path = os.path.join(audio_output_folder, f'{name}.wav')

            print("[INFO]: Type file:", type(audio.file))

            with tempfile.NamedTemporaryFile(delete=False) as temp_audio_file:
                shutil.copyfileobj(audio.file, temp_audio_file)
                temp_audio_file.close() 
                audio_clip = AudioFileClip(temp_audio_file.name)
                audio_clip.write_audiofile(audio_path, codec="pcm_s16le", ffmpeg_params=["-ac", "1"])
            os.remove(temp_audio_file.name)

        else:
            save_path = os.path.join(user_path, f'{name}{ext}')
            with open(save_path, "wb") as buffer:
                shutil.copyfileobj(audio.file, buffer)

        # Store both the UID and original filename for reference
        vid_infos.append({
            "uid": random_uid, 
            "chunked": False, 
            "session_id": session_id, 
            "filename": "".join([name, ".wav"])
            })

    if vid_infos:
        # mongo.add_source(vid_infos)
        supabase.add_or_update_section(session_id, "processing", "end_upload_audios_svc")
        supabase.add_source(vid_infos)
    return

#-------------------LLM PROCESSING---------------------------#
def fetch_summary(session_id: str):
    """
    Get LLM to summarize our script.

    Parameters:
    session_id (str): Unique identifier for this user.

    Returns:
    list[str]: Summarized results extracted from the text.
    """
    print(f"[DEBUG] fetch_summary: Starting for session {session_id}")
    # script_entry = mongo.find_script(session_id)
    script_entry = supabase.find_script(session_id)

    if not script_entry:
        print(f"[ERROR] fetch_summary: No script found for session ID: {session_id}")
        raise ValueError(f"No script found for session ID: {session_id}")
    
    script_uid = script_entry.get("uid")
    print(f"[DEBUG] fetch_summary: Found script with UID {script_uid}")

    trans_path = os.path.join(SERVICE_DIR, SCRIPT_DIR, session_id, script_uid + ".txt")
    # loader = UnstructuredFileLoader(file_path=trans_path)
    # loader = UnstructuredLoader(file_path=trans_path, chunking_strategy="basic", max_characters=1000000, include_orig_elements=False)
    # docs = loader.load()
    print(f"[DEBUG] fetch_summary: Script path: {trans_path}")

    # Check if file exists and is not empty
    if not os.path.exists(trans_path):
        print(f"[ERROR] fetch_summary: Script file does not exist: {trans_path}")
        raise FileNotFoundError(f"Script file not found: {trans_path}")
        
    # Read the file directly instead of using UnstructuredLoader
    try:
        print(f"[DEBUG] fetch_summary: Reading script file directly")
        with open(trans_path, 'r', encoding='utf-8') as file:
            script_content = file.read()
        print(f"[DEBUG] script content is : {script_content}")
        if not script_content.strip():
            print(f"[WARNING] fetch_summary: Script file is empty: {trans_path}")
            return [""]
            
        # loader = UnstructuredLoader(file_path=trans_path, chunking_strategy="basic", max_characters=1_000_000, include_orig_elements=False)
        docs = [LangchainDocument(page_content=script_content)]
        print("Number of LangChain documents:", len(docs))
        print("Length of text in the document:", len(docs[0].page_content))

        print(f"[DEBUG] fetch_summary: Successfully created document, length: {len(script_content)} chars")
        
        # Process with LLM
        print(f"[DEBUG] fetch_summary: Sending to LLM for summarization")
        output = stuff_chain.invoke({"text_input_here": docs})
        print(f"[DEBUG] fetch_summary: LLM output received: {output[:100]}...")
        
        task1_result = get_task1(output)
        print(f"[DEBUG] fetch_summary: Extracted Task 1 result with {len(task1_result)} items")
        return task1_result
    
    except Exception as e:
        print(f"[ERROR] fetch_summary: {str(e)}")
        return [""]
    
#----- fetch_summary_v2 ------#
def fetch_summary_v2(session_id: str, category: str) -> List[str]:
    """
    Get LLM to summarize script using a prompt from the given category.

    Parameters:
    - session_id (str): Unique identifier for this user.
    - category (str): type of prompt to use, e.g. "news", "

    Returns:
    - list[str]: Summarized results extracted from the text.
    """
    # script_entry = mongo.find_script(session_id)
    script_entry = supabase.find_script(session_id)
    if not script_entry:
        raise ValueError(f"No script found for session ID: {session_id}")
    
    script_uid = script_entry.get("uid")
    trans_path = os.path.join(SERVICE_DIR, SCRIPT_DIR, session_id, f"{script_uid}.txt")

    loader = UnstructuredLoader(
        file_path=trans_path,
        chunking_strategy="basic",
        max_characters=1_000_000,
        include_orig_elements=False
    )
    docs = loader.load()
    print("Number of LangChain documents:", len(docs))
    print("Length of text in the document:", len(docs[0].page_content))
    try:
        print(f"[INFO] fetch_summary_v2: Using category = '{category}'")

        # โหลด prompt จาก category
        prompt_template = get_prompt_by_category(category)

        # สร้าง chain ใหม่โดยใช้ openai_model เดิม
        local_chain = create_stuff_documents_chain(
            openai_model,
            prompt_template,
            document_variable_name="text_input_here"
        )
        output = local_chain.invoke({"text_input_here": docs})
        print("LLM Output:", output)
        task1_result = get_task1(output)
        return task1_result
    except Exception as e:
        print(f"[ERROR] fetch_summary_v2: {e}")
        return [""]

def get_task1(text) -> List[str]:
    """
    Extracts Task 1 information from the provided text using regular expressions.

    Parameters:
    text (str): Input text containing Task 1 prompt

    Returns:
    list[str]: List of strings containing Task 1 prompt
    """
    print("[INFO]: Start get task 1")
    reg = r"(?s)<Task_1>\s*(.*?)\s*<\/Task_1>"
    reg_result = re.findall(reg, text)
    reg_result = reg_result[0].replace("'", "").strip()
    reg_result = reg_result.split('\n')
    return reg_result

def summarize_svc(session_id: str, category: Optional[str] = None):
    """
    Summarizes a stored script and update status on DB.

    Parameters:
    session_id (str): Unique identifier for this user.
    category (Optional[str]): Prompt type to use ("normal", "news", etc.)

    Returns:
    dict: Contains the path to the summarized Task 1 file or error code
    """
    supabase.add_or_update_section(session_id, "processing", "start-summarize_svc")
    print(f"[DEBUG] summarize_svc: Starting for session {session_id}")

    # script_doc = mongo.find_script(session_id)
    script_doc = supabase.find_script(session_id)

    if not script_doc:
        print(f"[ERROR] summarize_svc: No script found for session ID: {session_id}")
        raise ValueError(f"No script found for session ID: {session_id}")
    
    # Save Task 1 summary to script_sum
    task1_path = os.path.join(SERVICE_DIR, SCRIPT_DIR, session_id,
                              f"{session_id}_task1.txt")
    print(f"[DEBUG] summarize_svc: Task1 path: {task1_path}")
    
    try:
        if not script_doc['summed']:
            # เพิ่ม optional category
            print(f"[DEBUG] summarize_svc: Script not yet summarized, fetching summary")
            task1_summary = fetch_summary(session_id)
            
            if not task1_summary or task1_summary == [""]:
                print(f"[WARNING] summarize_svc: Empty summary returned")
                return 1
                
            task1_summary = "\n".join(task1_summary)
            print(f"[DEBUG] summarize_svc: Writing summary to file, length: {len(task1_summary)} chars")
            
            # Ensure directory exists
            os.makedirs(os.path.dirname(task1_path), exist_ok=True)
            
            with open(task1_path, "w", encoding="utf-8") as file:
                file.write(task1_summary)

            print(f"[DEBUG] summarize_svc: Marking script as summarized in database")
            # mongo.done_llm(session_id)
            supabase.done_llm(session_id)
        else:
            print(f"[DEBUG] summarize_svc: Script already summarized, reading from file")
            if not os.path.exists(task1_path):
                print(f"[ERROR] summarize_svc: Summarized file not found: {task1_path}")
                return 1
                
            with open(task1_path, "r", encoding="utf-8") as file:
                task1_summary = file.read()

        # return {"task1": task1_path}
        supabase.add_or_update_section(session_id, "processing", "end-summarize_svc")

        print(f"[DEBUG] summarize_svc: Summary retrieval successful")
        return task1_summary
    
    except Exception as e:
        print(f"[ERROR] summarize_svc: {str(e)}")
        traceback.print_exc()
        return 1
    

#-------------------TTS GENERATION---------------------------#
def generate_audio(text: str, voice_id: str, session_id: str):
    """
    Generates audio from text using a Botnoi voice API.

    Parameters:
    text (str): Text to convert to audio.
    voice_id (str): ID for the speaker to use.
    session_id (str): Unique identifier for this user

    Returns:
    tuple: (audio_path, audio_url) if successful, raises an exception otherwise
    """
    print(f"[DEBUG] generate_audio: Starting TTS generation for session {session_id}")
    print(f"[DEBUG] generate_audio: Voice ID: {voice_id}")
    print(f"[DEBUG] generate_audio: Text length: {len(text)} characters")

    audio_dir = os.path.join(SERVICE_DIR, VOICE_DIR)
    
    # Ensure directory exists
    if not os.path.exists(audio_dir):
        os.makedirs(audio_dir, exist_ok=True)
        print(f"[DEBUG] generate_audio: Created audio directory: {audio_dir}")

    payload = payload_template.copy()
    payload["text"] = text
    payload["speaker"] = voice_id

    try:
        print(f"[DEBUG] generate_audio: Sending request to API with payload")
        response = requests.post(BOT_URL, headers=headers, json=payload, timeout=120)
        
        # Check if the request was successful
        if response.status_code != 200:
            print(f"[ERROR] generate_audio: API returned status code {response.status_code}")
            print(f"[ERROR] generate_audio: API response: {response.text}")
            raise Exception(f"API returned status code {response.status_code}: {response.text}")
        
        # Parse JSON response
        try:
            response_data = response.json()
            print(f"[DEBUG] generate_audio: API response: {response_data}")
        except Exception as e:
            print(f"[ERROR] generate_audio: Failed to parse API response as JSON: {str(e)}")
            raise Exception(f"Failed to parse API response as JSON: {str(e)}")

        if 'audio_url' in response_data:
            audio_url = response_data.get("audio_url")
            if not audio_url:
                print(f"[ERROR] generate_audio: 'audio_url' is empty in the response")
                raise Exception("Failed to retrieve audio URL: URL is empty")
            
            try:
                print(f"[DEBUG] generate_audio: Downloading audio from {audio_url}")
                audio = requests.get(audio_url, timeout=120)
                
                # Check if download was successful
                if audio.status_code != 200:
                    print(f"[ERROR] generate_audio: Failed to download audio: status {audio.status_code}")
                    raise Exception(f"Failed to download audio: status {audio.status_code}")
                
                final_path = os.path.join(audio_dir, f"{session_id}.wav")
                
                with open(final_path, 'wb') as audio_file:
                    audio_file.write(audio.content)
                    
                print(f"[DEBUG] generate_audio: Audio successfully saved at {final_path}")
                return final_path, audio_url
            except Exception as e:
                print(f"[ERROR] generate_audio: Failed to download and save audio: {str(e)}")
                raise Exception(f"Failed to download and save audio: {str(e)}")
        else:
            # Log all available keys to help debugging
            available_keys = list(response_data.keys()) if isinstance(response_data, dict) else "Response is not a dictionary"
            print(f"[ERROR] generate_audio: 'audio_url' key not found in response. Available keys: {available_keys}")
            
            # Sometimes the error message is in a different field
            error_message = response_data.get('message', 'Unknown error')
            raise Exception(f"Error in API response: {error_message}")

    except Exception as e:
        print(f"[ERROR] generate_audio: An error occurred: {str(e)}")
        traceback.print_exc()
        raise Exception(f"Failed to generate audio: {str(e)}")

# Change toggle2 keyword --> Do summarize the script OR The script is summarized
def gen_voice_svc(session_id: str, voice_id: str, original_script: bool = False):
    """
    Generates a voice from a script, either summarized or original.

    Parameters:
    session_id (str): Unique identifier for this user
    voice_id (str): ID for the speaker to use
    original_script (bool): Whether to use the original script (default False)

    Returns:
    tuple: (audio_path, audio_url) containing paths to the generated audio
    """
    supabase.add_or_update_section(session_id, "processing", "start-gen_voice_svc")
    print(f"[DEBUG] gen_voice_svc: Starting for session {session_id}, voice_id={voice_id}, original_script={original_script}")
    
    try:
        if not original_script:
            task1_path = os.path.join(SERVICE_DIR, SCRIPT_DIR, session_id, f"{session_id}_task1.txt")

            if not os.path.exists(task1_path):
                print(f"[ERROR] gen_voice_svc: Summarized script file not found: {task1_path}")
                raise FileNotFoundError(f"Summarized script file not found: {task1_path}")

            with open(task1_path, "r", encoding="utf-8") as file:
                lines = file.readlines()
                if not lines:
                    print(f"[ERROR] gen_voice_svc: Summarized script file is empty")
                    raise ValueError("Summarized script file is empty")
                    
                print(f"[DEBUG] gen_voice_svc: Read {len(lines)} lines from summarized script")
                # Process lines with delay
                combined_text = "delay{0.5}".join([line.strip()[:-1] if line.strip().endswith(".") else line.strip() for line in lines])
        else:
            # script_entry = mongo.find_script(session_id)
            script_entry = supabase.find_script(session_id)
            if not script_entry:
                print(f"[ERROR] gen_voice_svc: No script found for session ID: {session_id}")
                raise FileNotFoundError(f"No script found for session ID: {session_id}")

            script_uid = script_entry.get("uid")
            if not script_uid:
                print(f"[ERROR] gen_voice_svc: Script entry missing 'uid' for session ID: {session_id}")
                raise FileNotFoundError(f"Script entry missing 'uid' for session ID: {session_id}")

            script_path = os.path.join(SERVICE_DIR, SCRIPT_DIR, session_id, f"{script_uid}.txt")
            if not os.path.exists(script_path):
                print(f"[ERROR] gen_voice_svc: Original script file not found: {script_path}")
                raise FileNotFoundError(f"Original script file not found: {script_path}")

            with open(script_path, "r", encoding="utf-8") as file:
                combined_text = "delay{0.5}".join([line.strip()[:-1] if line.strip().endswith(".") else line.strip() for line in file.readlines()])

    except Exception as e:
        print(f"[ERROR] gen_voice_svc: {str(e)}")
        supabase.add_or_update_section(session_id, "fail", "error-gen_voice_svc")
        raise Exception(f"Failed to generate voice: {str(e)}")
    
    audio_path, audio_url = generate_audio(combined_text, voice_id, session_id)
    if audio_path is None:
        supabase.add_or_update_section(session_id, "fail", "error-gen_voice_svc")
        raise Exception("Failed to generate audio.")
    return audio_path, audio_url
    
def get_audio_length_svc(audio_path: str) -> str:
    """
    Retrieves the duration of an audio file.

    Parameters:
    audio_path (str): Path to the audio file

    Returns:
    str: Formatted duration in "MM:SS"
    """
    try:
        audioFile = AudioFileClip(audio_path)
        full = audioFile.duration
        audioFile.close()
        minutes = int(full / 60)
        seconds = int(full % 60)
        formatted = f"{minutes:02}:{seconds:02}"
        return formatted
    except Exception as e:
        print(f"Error getting audio duration: {e}")
        raise Exception(f"Failed to get audio duration: {audio_path}")

    
#-------------------TRANSCRIPTION SERVICE---------------------------#
def json_to_txt(session_id: str):
    """
    Converts JSON transcription data to a text (.txt) file.

    Parameters:
    session_id (str): Unique identifier for this user
    """
    json_path = os.path.join(SERVICE_DIR, CACHE_DIR, session_id, f'{session_id}.json')
    name = str(uuid4())
    txt_path = os.path.join(SERVICE_DIR, SCRIPT_DIR, session_id, f'{name}.txt')
    with open(json_path, encoding='utf-8') as json_file:
        raw = json.load(json_file)

    temp = raw.get('data')
    with open(txt_path, "w", encoding='utf-8') as file:
        for entry in temp:
            file.write(f'{entry[0]} - {entry[1]} {entry[2]}' + '\n')
            file.write('\n')
    
    file_and_owner = {"uid": name, 
                      "summed": False,
                      "session_id": session_id}
    
    # mongo.add_script(file_and_owner)
    supabase.add_script(file_and_owner)
    os.remove(json_path)

def transcribe_svc(session_id: str, diarize: bool = False):
    """
    Transcribes audio extracted from a video.

    Parameters:
    session_id (str): Unique identifier for the session
    diarize (bool): Flag for diarization in transcription (default False)

    Returns:
    str: Transcription result text
    """
    # videos = mongo.find_source(session_id)
    videos = supabase.find_source(session_id)
    
    if not videos:
        raise FileNotFoundError(f"No videos found for session ID: {session_id}")
    
    # Get the first video and its vuid
    first_video = videos[0]
    vuid = first_video["uid"]
    video_path = os.path.join(SERVICE_DIR, VID_SRC_DIR, session_id, f'{vuid}.mp4')
    
    if not os.path.exists(video_path):
        raise FileNotFoundError(f"Video file not found: {video_path}")
    
    # Extract audio from the video
    audio_path = os.path.join(SERVICE_DIR, CACHE_DIR, session_id, f'{vuid}.wav')
    extract_audio(video_path, audio_path)
    
    if not os.path.exists(audio_path):
        raise FileNotFoundError(f"Failed to extract audio from video: {video_path}")
    
    # Call the transcribe service with the extracted audio
    with open(audio_path, 'rb') as audio_file:
        files = {'audio': audio_file}
        params = {'diarize': str(diarize).lower()}
        
        try:
            response = requests.post(TRANSCRIBE_URL, files=files, params=params)
            response.raise_for_status()
        except requests.RequestException as e:
            raise HTTPException(status_code=500, detail=f"Transcription service error: {str(e)}")
    
    os.remove(audio_path)

    cache_file_path = os.path.join(SERVICE_DIR, CACHE_DIR, session_id, f"{session_id}.json")
    try:
        with open(cache_file_path, 'w', encoding='utf-8') as json_file:
            json.dump(response.json(), json_file, ensure_ascii=False, indent=4)
    except Exception as e:
        raise Exception(f"Failed to save transcription response as JSON: {str(e)}")

    json_to_txt(session_id)

    return response.text

def extract_audio(video_path: str, audio_path: str):
    """
    Extracts audio from a video file and saves it as a WAV file.

    Parameters:
    video_path (str): Path to the video file
    audio_path (str): Path where the extracted audio should be saved
    """
    try:
        if not os.path.exists(os.path.dirname(audio_path)):
            os.makedirs(os.path.dirname(audio_path))
        video = VideoFileClip(video_path) 
        audio = video.audio  
        audio.write_audiofile(audio_path, codec="pcm_u8")  # Save the audio as a .wav file
        audio.close()
        video.close() 
    except Exception as e:
        print(f"Error extracting audio: {str(e)}")
        raise Exception(f"Failed to extract audio from video: {video_path}")
    
#-------------------30 sec-----------------------------------------------#
def extract_audio_v2(video_path: str, filename: str, session_id: str):
    """
    Extracts and splits audio into 30-second segments.

    Parameters:
    video_path (str): Path to the video file
    session_id (str): Unique identifier for this user
    """
    try:
        audio_output_folder = os.path.join(SERVICE_DIR, CACHE_DIR, session_id, "audio_segments")
        if not os.path.exists(audio_output_folder):
            os.makedirs(audio_output_folder, exist_ok=True)
        name, ext = os.path.splitext(filename)
        audio_path = os.path.join(audio_output_folder, f'{name}.wav')
        
        if not os.path.exists(audio_path):

            video = VideoFileClip(video_path)
            video.audio.write_audiofile(audio_path, codec="pcm_s16le", ffmpeg_params=["-ac", "1"])  #pcm_s16le , pcm_u8
            video.close() 

        if len(os.listdir(audio_output_folder)) < 2:
            audio = AudioSegment.from_wav(audio_path)
            total_duration = len(audio) / 1000  
            num_segments = math.ceil(total_duration / SEGMENT_LENGTH)

            # Split audio into 30-second segments
            for i in range(num_segments):
                start_time = i * SEGMENT_LENGTH * 1000 
                end_time = min((i + 1) * SEGMENT_LENGTH * 1000, len(audio)) 
                segment = audio[start_time:end_time]
                segment_path = os.path.join(audio_output_folder, f'{session_id}_segment_{i + 1}.wav')
                segment.export(segment_path, format="wav")
                print(f"Exported segment: {segment_path}")

        # Remove the full audio file
        os.remove(audio_path)
        print(f"Removed full audio file: {audio_path}")

    except Exception as e:
        print(f"Error extracting audio: {str(e)}")
        raise Exception(f"Failed to extract audio from video: {video_path}")

def asr(audio_file, lang='th'):
    """
    Transcribes audio to text using ASR (Automatic Speech Recognition).

    Parameters:
    audio_file (str): Path to the audio file
    lang (str): Language code for transcription (default 'th')

    Returns:
    str: Transcribed text
    """
    print("audio_file:", audio_file)
    with open(audio_file, 'rb') as f:
        audio_data = f.read()
        
        payload =json.dumps({
                    "audio": base64.encodebytes(audio_data).decode('utf-8'),
                    "phone":'66999999998',
                    "speaker_lang":lang,
                    "model_name": MODEL_ASR, # whisper distil-whisper-ft-v2
                    "chunk_duration": SEGMENT_LENGTH
        })
        # Make API request
        response = requests.post(TRANSCRIBE2_URL, data=payload)
        text = response.text
        print("response:", text)
    return (text) 

async def transcribe_svc_v2(session_id: str, lang: str = 'th', diarize: bool = False):
    """
    Transcribes audio by splitting it into segments and processing asynchronously.

    Parameters:
    session_id (str): Unique identifier for the session
    lang (str): Language code for transcription (default 'th')
    diarize (bool): Flag for diarization in transcription (default False)

    Returns:
    dict: Transcript data as a dictionary
    """
    try:
        import logging
        logger = logging.getLogger(__name__)
        logger.warning("Starting transcribe")

        supabase.add_or_update_section(session_id, "processing", "start-transcribe_svc_v2")
        
        # Find the video for the session
        # videos = mongo.find_source(session_id)
        videos = supabase.find_source(session_id)
        if not videos:
            raise FileNotFoundError(f"No videos found for session ID: {session_id}")
        logger = logging.getLogger(__name__)
        logger.warning(videos)

        # Get the first video and its vuid
        first_video = videos[0]
        filename = first_video["filename"]
        video_path = os.path.join(SERVICE_DIR, VID_SRC_DIR, session_id, f'{filename}')
        logger.warning("///////////////////////////////////////////////////////////")
        logger.warning(video_path)
        logger.warning("///////////////////////////////////////////////////////////")
        audio_folder = os.path.join(SERVICE_DIR, CACHE_DIR, session_id, "audio_segments")
        audio_path = os.path.join(audio_folder, f'{filename}')

        if not os.path.exists(video_path) and not os.path.exists(audio_path):
            raise FileNotFoundError(f"Video/Audio file not found: {video_path} or {audio_path}")

        # Step 1: Extract and split audio into segments
        extract_audio_v2(video_path, filename, session_id)

        # Step 2: Send each segment for transcription
        # audio_folder = os.path.join(SERVICE_DIR, CACHE_DIR, session_id, "audio_segments")
        transcription_folder = os.path.join(SERVICE_DIR, CACHE_DIR, session_id, "transcriptions")
        os.makedirs(transcription_folder, exist_ok=True)

        audio_files = sorted([f for f in os.listdir(audio_folder) if f.endswith('.wav')],
                     key=lambda x: int(x.split('_')[-1].split('.')[0]) if x.split('_')[-1].split('.')[0].isdigit() else 0)

        params = {'diarize': str(diarize).lower()}

        combined_transcript = ""
        import time
        start_asr = time.time()
        task = []
        for audio_file in audio_files:
            audio_path = os.path.join(audio_folder, audio_file)
            task.append(asyncio.create_task(asyncio.to_thread(asr, audio_path, lang)))

        output = await asyncio.gather(*task)
        print("[ASR TIME]: ", time.time() - start_asr)

        count_output = 0
        for i, transcription in enumerate(output, start=1):
            try:
                transcription = json.loads(transcription)
                transcription_str = '\n' + transcription['result']['text'] + '\n'

                combined_transcript += transcription_str + "\n"

                transcription_path = os.path.join(transcription_folder, f"{audio_files[i-1].replace('.wav', '.txt')}")
                with open(transcription_path, 'w', encoding='utf-8') as transcription_file:
                    transcription_file.write(transcription_str)

                print(f"Saved transcription for segment {i}: {transcription_path}")
            except Exception as e:
                count_output += 1
                if count_output > int(len(output)*0.8):
                    print(f"Error transcribe audio: {str(e)}")
                    raise Exception(f"failed to transcribe audio from video: {video_path}")

            # with open(audio_path, 'rb') as audio_data:
            #     files = {'audio': audio_data}
            #     try:
            #         response = requests.post(TRANSCRIBE_URL, files=files, params=params)
            #         response.raise_for_status()

            #         # Get the transcription result
            #         transcription = response.json().get('data', '')

            #         # Handle cases where transcription is a list
            #         transcription_str = '\n'.join([f"{item[0]} - {item[1]}: {item[2]}" for item in transcription]) if isinstance(transcription, list) else transcription

            #         # Add each segment's transcription to the combined transcript
            #         combined_transcript += transcription_str + "\n"

            #         # Save transcription to a text file 
            #         transcription_path = os.path.join(transcription_folder, f"{audio_file.replace('.wav', '.txt')}")
            #         with open(transcription_path, 'w', encoding='utf-8') as transcription_file:
            #             transcription_file.write(transcription_str)

            #         print(f"Saved transcription for segment {i}: {transcription_path}")
            #     except requests.RequestException as e:
            #         raise Exception(f"Failed to transcribe segment {i}: {str(e)}")

        # Step 3: Combine transcriptions into a single file
        combine_transcriptions_v2(transcription_folder, session_id)

        transcript_data = get_transcript_data(session_id)

        # Step 4: Clean up (remove the session folder)
        shutil.rmtree(audio_folder)
        shutil.rmtree(transcription_folder)

        print(f"Cleaned up session folder: {audio_folder}")
        print(f"Cleaned up session folder: {transcription_folder}")

        # Add transcribed text to database with session_id
        # transcribe_text = ' '.join(transcript_data['data'])
        # supabase.add_script({
        #     'uid': first_video["uid"],
        #     'session_id': session_id,
        #     'transcribe_text': transcribe_text,
        #     'summed': False
        # })
        
        supabase.add_or_update_section(session_id, "processing", "end-transcribe_svc_v2")
        return transcript_data 
    
    except Exception as e:
        print(f"Error in transcription service: {str(e)}")
        supabase.add_or_update_section(session_id, "fail", "error-transcribe_svc_v2")
        
        raise Exception(f"Transcription service failed for session {session_id}")


def combine_transcriptions_v2(transcription_folder: str, session_id: str):
    """
    Combines individual transcriptions into one file with correct timestamps.

    Parameters:
    transcription_folder (str): Folder containing transcription files 
    session_id (str): Unique identifier for this user
    """
    
    try:
        # Sort transcription files by segment number
        transcription_files = sorted(
            [f for f in os.listdir(transcription_folder) if f.endswith('.txt')],
            key=lambda x: int(re.search(r'segment_(\d+)', x).group(1)) if re.search(r'segment_(\d+)', x) else 0
        )

        output_file_path = os.path.join(SERVICE_DIR, SCRIPT_DIR, session_id, f'{session_id}.txt')
        if not os.path.exists(os.path.dirname(output_file_path)):
            os.makedirs(os.path.dirname(output_file_path), exist_ok=True)

        # Initialize current_time with 0 (which would be the starting time of segment 1)
        current_time = 0
        
        # Store all transcribed text
        all_text = []

        # Combine all transcriptions into one file with correct segment timestamps
        with open(output_file_path, 'w', encoding='utf-8') as output_file:
            for transcription_file in transcription_files:
                transcription_path = os.path.join(transcription_folder, transcription_file)

                # Extract segment number from the file name
                match = re.search(r'segment_(\d+)', transcription_file)
                if match:
                    segment_number = int(match.group(1))

                    # Calculate the expected start time based on segment number
                    expected_start_time = (segment_number - 1) * SEGMENT_LENGTH

                    # Ensure the current_time matches the expected start time for this segment
                    if current_time < expected_start_time:
                        # If there is a gap, insert empty lines with the correct timestamps
                        while current_time < expected_start_time:
                            new_start_time = convert_seconds_to_time(current_time)
                            new_end_time = convert_seconds_to_time(current_time + SEGMENT_LENGTH)
                            output_file.write(f"{new_start_time} - {new_end_time}: [No transcription available]\n")
                            all_text.append("[No transcription available]")
                            current_time += SEGMENT_LENGTH

                with open(transcription_path, 'r', encoding='utf-8') as file:
                    lines = file.readlines()

                    for line in lines:
                        # Extract timestamp and text
                        match = re.match(r'(\d+:\d+:\d+) - (\d+:\d+:\d+): (.*)', line)
                        if match:
                            start_time, end_time, text = match.groups()
                            all_text.append(text.strip())

                            # Convert start and end times to seconds
                            start_seconds = convert_time_to_seconds(start_time)
                            end_seconds = convert_time_to_seconds(end_time)

                            # Adjust timestamps to align with the current cumulative time
                            new_start_time = convert_seconds_to_time(current_time)
                            current_time += (end_seconds - start_seconds)
                            new_end_time = convert_seconds_to_time(current_time)

                            # Write the adjusted line to the output file
                            output_file.write(f"{new_start_time} - {new_end_time}: {text}\n")
                        else:
                            output_file.write(line)
                            all_text.append(line.strip())

                # After processing the current file, advance current_time by SEGMENT_LENGTH for the next segment
                current_time = (segment_number) * SEGMENT_LENGTH

        # Combine all text and add to metadata
        transcribe_text = ' '.join([text for text in all_text if text]) 
        metadata = {
            "uid": session_id,
            "summed": False, 
            "session_id": session_id,
            "transcribe_text": transcribe_text
        }
        
        # mongo.add_script(metadata)
        supabase.add_script(metadata)
        print(f"Combined transcription saved at: {output_file_path}")

    except Exception as e:
        print(f"Error combining transcriptions: {str(e)}")
        raise Exception(f"Failed to combine transcriptions for session {session_id}")


def convert_time_to_seconds(time_str):
    """
    Converts a time string in "HH:MM:SS" format to the total number of seconds.

    Parameters:
    time_str (str): A string representing time in "HH:MM:SS" format.

    Returns:
    int: The total time converted to seconds.
    """
    h, m, s = map(int, time_str.split(':'))
    return h * 3600 + m * 60 + s


def convert_seconds_to_time(seconds):
    """
    Converts a total number of seconds to a time string in "HH:MM:SS" format.

    Parameters:
    seconds (int): The total number of seconds to convert.

    Returns:
    str: A string representing the time in "HH:MM:SS" format.
    """
    h = seconds // 3600
    m = (seconds % 3600) // 60
    s = seconds % 60
    return f"{h:02}:{m:02}:{s:02}"

def get_transcript_data(session_id: str):
    """
    Retrieves and parses the transcript data for a given session ID from a text file.

    Parameters:
    session_id (str): The unique identifier of the session whose transcript is to be retrieved.

    Returns:
    dict: A dictionary containing parsed transcript data with "data" as the key, where each entry is a list containing 
          start time, end time, and text for each transcript segment.
    """
    try:
        output_file_path = os.path.join(SERVICE_DIR, SCRIPT_DIR, session_id, f'{session_id}.txt')
        data = []
        with open(output_file_path, 'r', encoding='utf-8') as f:
            for line in f:
                match = re.match(r'(\d{2}:\d{2}:\d{2}) - (\d{2}:\d{2}:\d{2}):\s*(.*)', line)
                if match:
                    start_time, end_time, text = match.groups()
                    data.append([start_time, end_time, text.strip()])
                else:
                    data.append(line.strip())
        return {"data": data}
    except Exception as e:
        print(f"Error parsing transcription file: {str(e)}")
        raise Exception(f"Failed to parse transcription file for session {session_id}")

